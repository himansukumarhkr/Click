import os
import threading
import ctypes
import tempfile
import shutil
import queue
import time
import io
import re
from typing import Optional, List, Dict

from PIL import Image, ImageGrab
from docx import Document
from docx.shared import Inches
import win32clipboard
import win32con

from src.hotkeys import kernel32, user32


class ScreenshotSession:
    """
    Manages a single screenshot session.
    Handles capturing images, saving to Word/Folder, and clipboard operations.
    """

    def __init__(self, config: Dict, gui_callback_queue: queue.Queue):
        self.config = config
        self.gui_queue = gui_callback_queue

        self.base_filename = ""
        self.current_filepath = ""
        self.screenshot_count = 0
        self.max_size_bytes = 0
        self.captured_images = []
        self.temp_dir = None
        self.document = None
        self.is_running = True
        self.status = "Active"

        self.save_queue = queue.Queue()
        self.clipboard_queue = queue.Queue()
        self.save_lock = threading.Lock()

        self.save_thread = None
        self.clipboard_thread = None

        self.warning_shown = False
        self.last_size_str = "0 KB"

        self._initialize_session()
        self.session_id = self.current_filepath

    def _initialize_session(self):
        """Sets up the save directory and initial file."""
        save_dir = self.config['save_dir']
        filename_input = self.config['filename']

        if not os.path.exists(save_dir):
            try:
                os.makedirs(save_dir)
            except OSError:
                pass

        full_path = os.path.join(save_dir, filename_input)

        if self.config['save_mode'] == 'folder':
            if not os.path.exists(full_path):
                self.current_filepath = full_path
            else:
                self.current_filepath = self._get_unique_path(full_path)

            self.base_filename = os.path.basename(self.current_filepath)
            if not os.path.exists(self.current_filepath):
                os.makedirs(self.current_filepath)
        else:
            self.current_filepath, self.base_filename = self._get_unique_file(save_dir, filename_input)
            self.document = Document()
            try:
                self.document.save(self.current_filepath)
            except OSError:
                pass

        try:
            self.max_size_bytes = int(float(self.config['max_size']) * 1024 * 1024)
        except (ValueError, TypeError):
            self.max_size_bytes = 0

        self.temp_dir = tempfile.mkdtemp(prefix="Click_")
        self._start_workers()

    def _get_unique_path(self, path: str) -> str:
        counter = 1
        while True:
            new_path = f"{path}_{counter}"
            if not os.path.exists(new_path):
                return new_path
            counter += 1

    def _get_unique_file(self, directory: str, name: str):
        counter = 0
        while True:
            new_name = name if counter == 0 else f"{name}_{counter}"
            file_path = os.path.join(directory, new_name + ".docx")
            if not os.path.exists(file_path):
                return file_path, new_name
            counter += 1

    def _start_workers(self):
        if not self.save_thread or not self.save_thread.is_alive():
            self.save_thread = threading.Thread(target=self._save_worker, daemon=True)
            self.save_thread.start()

        if not self.clipboard_thread or not self.clipboard_thread.is_alive():
            self.clipboard_thread = threading.Thread(target=self._clipboard_worker, daemon=True)
            self.clipboard_thread.start()

    def stop(self):
        self.is_running = False
        if self.document:
            try:
                self.document.save(self.current_filepath)
            except OSError:
                pass

    def capture(self):
        if not self.is_running:
            return

        try:
            image = ImageGrab.grab(all_screens=True)
        except Exception:
            return

        self.screenshot_count += 1
        window_title = self._get_active_window_title() if self.config['log_title'] else None

        self.gui_queue.put(("NOTIFY", self.session_id, self.screenshot_count, self.last_size_str))

        if self.config['auto_copy']:
            if self.config['save_mode'] == "folder":
                fname = f"{self.base_filename}_{self.screenshot_count}.jpg"
            else:
                fname = f"screen_{self.screenshot_count}.jpg"

            temp_path = os.path.join(self.temp_dir, fname)

            # We need the image data if we are copying the image directly OR if we need to save the file for file-copy
            should_pass_image = self.config.get('copy_image', True) or self.config.get('copy_files', True)
            copy_img_data = image if should_pass_image else None

            # Cumulative auto-copy: include all previous files + current one
            all_files = list(self.captured_images) + [temp_path]
            self.clipboard_queue.put((copy_img_data, temp_path, all_files))

        self.save_queue.put((image, self.screenshot_count, window_title))

    def undo(self):
        if self.is_running:
            self.save_queue.put(("UNDO", None, None))

    def manual_rotate(self):
        if self.config['save_mode'] != "folder":
            self.save_queue.put(("ROTATE", None, None))

    def _save_worker(self):
        while True:
            try:
                try:
                    task = self.save_queue.get(timeout=0.1)
                except queue.Empty:
                    if not self.is_running and self.save_queue.empty():
                        break
                    continue

                if task[0] == "UNDO":
                    self._perform_undo()
                elif task[0] == "ROTATE":
                    self._rotate_file()
                else:
                    self._perform_save(*task)

                self.save_queue.task_done()
            except Exception as e:
                print(f"Worker Error: {e}")

    def _perform_save(self, image, count, window_title):
        try:
            if self.config['save_mode'] == "folder":
                filename = f"{self.base_filename}_{count}.jpg"
            else:
                filename = f"screen_{count}.jpg"

            image_path = os.path.join(self.temp_dir, filename)

            with self.save_lock:
                if not os.path.exists(image_path):
                    image.save(image_path, "JPEG", quality=90)

            self.captured_images.append(image_path)

            if self.config['save_mode'] == "folder":
                shutil.copyfile(image_path, os.path.join(self.current_filepath, filename))
                self.last_size_str = self._get_folder_size(self.current_filepath)
            else:
                if os.path.exists(self.current_filepath) and self.max_size_bytes > 0:
                    current_size = os.path.getsize(self.current_filepath)
                    if (current_size + os.path.getsize(image_path)) > self.max_size_bytes:
                        self._rotate_file()

                if not self.document:
                    self.document = Document(self.current_filepath)

                caption = ""
                if window_title and self.config['append_num']:
                    caption = f"{window_title} {count}"
                elif window_title:
                    caption = window_title
                elif self.config['append_num']:
                    caption = str(count)

                if caption:
                    self.document.add_paragraph(caption)

                self.document.add_picture(image_path, width=Inches(6))
                self.document.add_paragraph("-" * 50)

                try:
                    self.document.save(self.current_filepath)
                    self.last_size_str = self._get_file_size(self.current_filepath)
                    self.warning_shown = False
                except PermissionError:
                    if not self.warning_shown:
                        self.gui_queue.put(
                            ("WARNING", "File Locked", "Please close the Word document to continue saving."))
                        self.warning_shown = True
                    self.last_size_str = "File Locked"

            # Use session_id for UI updates to ensure the main thread can find the correct session
            self.gui_queue.put(("UPDATE_SESSION", self.session_id, self.screenshot_count, self.last_size_str))
        except Exception as e:
            print(f"Save Error: {e}")

    def _perform_undo(self):
        if self.screenshot_count <= 0:
            return

        try:
            if self.captured_images:
                path = self.captured_images.pop()
                if os.path.exists(path):
                    os.remove(path)

            if self.config['save_mode'] == 'folder':
                file_to_remove = os.path.join(self.current_filepath,
                                              f"{self.base_filename}_{self.screenshot_count}.jpg")
                if os.path.exists(file_to_remove):
                    os.remove(file_to_remove)
                self.last_size_str = self._get_folder_size(self.current_filepath)

            elif self.document and len(self.document.paragraphs) >= 2:
                try:
                    removed_count = 0
                    while removed_count < 3 and self.document.paragraphs:
                        p = self.document.paragraphs[-1]
                        p._element.getparent().remove(p._element)
                        removed_count += 1
                        if removed_count == 1 and "-" not in p.text and len(p.text) > 0:
                            break

                    self.document.save(self.current_filepath)
                    self.last_size_str = self._get_file_size(self.current_filepath)
                except Exception:
                    pass

            self.screenshot_count -= 1
            self.gui_queue.put(("UNDO", self.session_id, self.screenshot_count, self.last_size_str))
        except Exception:
            pass

    def cleanup(self, delete_files=False):
        self.stop()
        if self.temp_dir and os.path.exists(self.temp_dir):
            try:
                shutil.rmtree(self.temp_dir)
            except OSError:
                pass

        if delete_files and self.current_filepath:
            try:
                if self.config['save_mode'] == 'folder':
                    if os.path.exists(self.current_filepath):
                        shutil.rmtree(self.current_filepath)
                else:
                    # Delete all split parts
                    directory = os.path.dirname(self.current_filepath)
                    base = os.path.basename(self.current_filepath)

                    # Try to find the root name (e.g., Session from Session_Part2.docx)
                    match = re.search(r"^(.*)_Part\d+\.docx$", base)
                    if match:
                        root_name = match.group(1)
                    else:
                        root_name = base.rsplit('.', 1)[0]

                    # Delete the base file if it exists (e.g. Session.docx)
                    base_path = os.path.join(directory, f"{root_name}.docx")
                    if os.path.exists(base_path):
                        try:
                            os.remove(base_path)
                        except OSError:
                            pass

                    # Delete all parts (Session_Part1.docx, Session_Part2.docx, ...)
                    if os.path.exists(directory):
                        for file in os.listdir(directory):
                            if file.startswith(f"{root_name}_Part") and file.endswith(".docx"):
                                try:
                                    os.remove(os.path.join(directory, file))
                                except OSError:
                                    pass
            except OSError:
                pass

    def _rotate_file(self):
        if self.document:
            try:
                self.document.save(self.current_filepath)
            except OSError:
                pass
        self.document = None

        directory = os.path.dirname(self.current_filepath)
        base = os.path.basename(self.current_filepath).rsplit('.', 1)[0]

        match = re.search(r"^(.*)_Part(\d+)$", base)

        if match:
            root_name = match.group(1)
            current_num = int(match.group(2))
            new_name = f"{root_name}_Part{current_num + 1}.docx"
            self.current_filepath = os.path.join(directory, new_name)
        else:
            root_name = base
            counter = 1
            while True:
                part1_name = f"{root_name}_Part{counter}.docx"
                part1_path = os.path.join(directory, part1_name)
                if not os.path.exists(part1_path):
                    break
                counter += 1

            if os.path.exists(self.current_filepath):
                try:
                    os.rename(self.current_filepath, part1_path)
                except OSError:
                    pass

            new_name = f"{root_name}_Part{counter + 1}.docx"
            self.current_filepath = os.path.join(directory, new_name)

        self.document = Document()
        self.document.save(self.current_filepath)
        self.last_size_str = "0 KB"
        self.gui_queue.put(("UPDATE_FILENAME", self.session_id, self.current_filepath))

    def _get_active_window_title(self):
        try:
            hwnd = user32.GetForegroundWindow()
            length = user32.GetWindowTextLengthW(hwnd)
            buff = ctypes.create_unicode_buffer(length + 1)
            user32.GetWindowTextW(hwnd, buff, length + 1)
            title = buff.value
            return title.replace(" - Google Chrome", "").replace(" - Microsoft Edge", "")
        except Exception:
            return "Unknown"

    def _get_file_size(self, path):
        if not os.path.exists(path):
            return "0 KB"
        size = os.path.getsize(path)
        return f"{size / 1024:.2f} KB" if size < 1048576 else f"{size / 1048576:.2f} MB"

    def _get_folder_size(self, path):
        total = sum(os.path.getsize(os.path.join(r, f)) for r, _, fs in os.walk(path) for f in fs)
        return f"{total / 1024:.2f} KB" if total < 1048576 else f"{total / 1048576:.2f} MB"

    def _clipboard_worker(self):
        while True:
            try:
                try:
                    item = self.clipboard_queue.get(timeout=0.1)
                except queue.Empty:
                    if not self.is_running:
                        break
                    continue

                # Unpack arguments (handle legacy 2-item tuple if needed, though we updated capture)
                if len(item) == 3:
                    image_data, save_path, clipboard_files = item
                else:
                    image_data, save_path = item
                    clipboard_files = [save_path]

                # Ensure the file exists for file-copy (save current capture if needed)
                if image_data:
                    with self.save_lock:
                        if not os.path.exists(save_path):
                            image_data.convert("RGB").save(save_path, "JPEG", quality=90)
                            time.sleep(0.05)  # Ensure file system is ready

                # Copy to clipboard (using the cumulative list 'clipboard_files')
                self.copy_to_clipboard(image_data, clipboard_files)
                self.clipboard_queue.task_done()
            except Exception as e:
                print(f"Clipboard Error: {e}")

    def copy_to_clipboard(self, image, file_paths):
        copy_img = self.config.get('copy_image', True)
        copy_files = self.config.get('copy_files', True)

        if not copy_img and not copy_files:
            return

        try:
            # Retry opening clipboard a few times
            for _ in range(5):
                try:
                    win32clipboard.OpenClipboard()
                    break
                except Exception:
                    time.sleep(0.1)
            else:
                print("Failed to open clipboard after retries")
                return

            try:
                win32clipboard.EmptyClipboard()

                # 1. Set File List (CF_HDROP) - Priority for Explorer
                if copy_files and file_paths:
                    files_text = "\0".join([os.path.abspath(p) for p in file_paths]) + "\0\0"
                    files_data = files_text.encode('utf-16le')
                    # DROPFILES structure: pFiles(4), pt(8), fNC(4), fWide(4)
                    header = b'\x14\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x01\x00\x00\x00'
                    drop_data = header + files_data
                    win32clipboard.SetClipboardData(win32clipboard.CF_HDROP, drop_data)

                # 2. Set Image (CF_DIB) - For Visual History/Paint
                if copy_img and image:
                    output = io.BytesIO()
                    image.convert("RGB").save(output, "BMP")
                    data = output.getvalue()[14:]
                    output.close()
                    win32clipboard.SetClipboardData(win32clipboard.CF_DIB, data)

                # 3. Set Text Fallback (CF_UNICODETEXT) - Only if no image
                # Add text representation so it appears in Win+V history
                # This also allows pasting file paths as text in editors
                # Only add text if we are NOT copying an image, otherwise Win+V might show text instead of image
                if copy_files and file_paths and not (copy_img and image):
                    display_text = "\r\n".join([os.path.abspath(p) for p in file_paths])
                    try:
                        win32clipboard.SetClipboardData(win32clipboard.CF_UNICODETEXT, display_text)
                    except Exception:
                        pass  # Non-critical if text fallback fails

            finally:
                win32clipboard.CloseClipboard()

        except Exception as e:
            print(f"Clipboard Error: {e}")
            import traceback
            traceback.print_exc()

    def manual_copy_all(self):
        # Start a thread to process the copy loop with delays
        threading.Thread(target=self._process_manual_copy_all, daemon=True).start()

    def _process_manual_copy_all(self):
        if not self.captured_images:
            return

        total = len(self.captured_images)

        # 1. Populate History: Copy previous images one by one with delay
        # Skip the last one for now, as it will be handled in the final step
        images_to_process = self.captured_images[:-1]

        # Only do history population if copy_image is enabled
        if self.config.get('copy_image', True):
            for i, img_path in enumerate(images_to_process):
                # Send Progress (1-based index)
                self.gui_queue.put(("COPY_PROGRESS", i + 1, total))

                try:
                    if os.path.exists(img_path):
                        with Image.open(img_path) as img:
                            # Copy image only to populate history
                            # We don't need to copy files here, just the bitmap for visual history
                            self.copy_to_clipboard(img.copy(), [])
                        time.sleep(1.0)  # 1 second delay as requested
                except Exception:
                    pass

        # 2. Final Step: Copy Last Image + ALL Files
        # This ensures the final clipboard state allows pasting files into Explorer
        # and shows the most recent image in history/bitmap paste
        last_image = None

        # Send Final Progress
        self.gui_queue.put(("COPY_PROGRESS", total, total))

        if self.config.get('copy_image', True):
            try:
                last_path = self.captured_images[-1]
                if os.path.exists(last_path):
                    with Image.open(last_path) as img:
                        last_image = img.copy()
            except Exception:
                pass

        self.copy_to_clipboard(last_image, self.captured_images)

    def copy_master_file_to_clipboard(self):
        if self.document:
            try:
                self.document.save(self.current_filepath)
            except OSError:
                pass
        self.copy_to_clipboard(None, [os.path.abspath(self.current_filepath)])