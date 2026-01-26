import keyboard
from PIL import Image, ImageGrab
from docx import Document
from docx.shared import Inches
import os
import datetime
import tkinter as tk
from tkinter import filedialog, ttk
import threading
import ctypes
from ctypes import wintypes
import tempfile
import shutil
import queue
import time
import io

kernel32 = ctypes.windll.kernel32
user32 = ctypes.windll.user32

SIZE_T = ctypes.c_size_t
HGLOBAL = wintypes.HGLOBAL
LPVOID = ctypes.c_void_p
BOOL = wintypes.BOOL
UINT = wintypes.UINT
HANDLE = wintypes.HANDLE
HWND = wintypes.HWND

GMEM_FIXED = 0x0000
GMEM_ZEROINIT = 0x0040
GPTR = GMEM_FIXED | GMEM_ZEROINIT

kernel32.GlobalAlloc.argtypes = [UINT, SIZE_T]
kernel32.GlobalAlloc.restype = HGLOBAL

kernel32.GlobalLock.argtypes = [HGLOBAL]
kernel32.GlobalLock.restype = LPVOID

kernel32.GlobalUnlock.argtypes = [HGLOBAL]
kernel32.GlobalUnlock.restype = BOOL

kernel32.GlobalFree.argtypes = [HGLOBAL]
kernel32.GlobalFree.restype = HGLOBAL

user32.OpenClipboard.argtypes = [HWND]
user32.OpenClipboard.restype = BOOL

user32.EmptyClipboard.argtypes = []
user32.EmptyClipboard.restype = BOOL

user32.SetClipboardData.argtypes = [UINT, HANDLE]
user32.SetClipboardData.restype = HANDLE

user32.CloseClipboard.argtypes = []
user32.CloseClipboard.restype = BOOL


class POINT(ctypes.Structure):
    _fields_ = [("x", ctypes.c_long), ("y", ctypes.c_long)]


class DROPFILES(ctypes.Structure):
    _fields_ = [
        ("pFiles", wintypes.DWORD),
        ("pt", POINT),
        ("fNC", wintypes.BOOL),
        ("fWide", wintypes.BOOL),
    ]


class ScreenshotSession:
    def __init__(self):
        self.base_name_no_ext = "screenshots"
        self.current_filename = None
        self.session_count = 0
        self.max_size_bytes = 0
        self.current_part = 1
        self.is_split_mode = False
        self.log_window_titles = False
        self.append_sequence_number = True
        self.auto_copy_clipboard = False
        self.copy_files_option = True
        self.copy_image_option = True
        self.save_mode = "docx"
        self.running = False
        self.save_directory = ""
        self.image_paths = []
        self.temp_dir = None
        self.doc = None

        # Threading & Queues
        self.save_queue = queue.Queue()
        self.gui_callback = None
        self.worker_thread = None
        self.file_locked_warning_shown = False
        self.last_known_size_str = "0 KB"

    def set_callback(self, callback):
        self.gui_callback = callback

    def start_worker(self):
        if self.worker_thread is None or not self.worker_thread.is_alive():
            self.worker_thread = threading.Thread(target=self._worker_loop, daemon=True)
            self.worker_thread.start()

    def _worker_loop(self):
        while True:
            try:
                try:
                    task = self.save_queue.get(timeout=0.1)
                except queue.Empty:
                    if not self.running and self.save_queue.empty():
                        break
                    continue

                if task is None: break

                if task[0] == "UNDO":
                    self._process_undo()
                else:
                    img, count, window_title = task
                    self._process_save(img, count, window_title)

                self.save_queue.task_done()
            except Exception as e:
                print(f"Worker Error: {e}")

    def _process_save(self, img, count, window_title):
        try:
            if not self.temp_dir or not os.path.exists(self.temp_dir):
                self.temp_dir = tempfile.mkdtemp(prefix="ScreenshotTool_")

            if self.save_mode == "folder":
                img_filename = f"{self.base_name_no_ext}_{count}.jpg"
            else:
                img_filename = f"screenshot_{count}.jpg"
                
            img_path = os.path.join(self.temp_dir, img_filename)

            img.save(img_path, "JPEG", quality=90)
            self.image_paths.append(img_path)

            if self.auto_copy_clipboard:
                self.copy_dual_to_clipboard(img, [img_path])

            if self.save_mode == "folder":
                # Ensure the folder exists
                if not os.path.exists(self.current_filename):
                    os.makedirs(self.current_filename)
                    
                target_path = os.path.join(self.current_filename, img_filename)
                try:
                    for _ in range(3):
                        try:
                            shutil.copyfile(img_path, target_path)
                            break
                        except PermissionError:
                            time.sleep(0.1)
                    else:
                        shutil.copyfile(img_path, target_path)

                    self.last_known_size_str = self.get_folder_size_str(self.current_filename)
                    if self.gui_callback:
                        self.gui_callback("UPDATE_SIZE", self.last_known_size_str)
                except Exception as e:
                    print(f"Folder Save Error: {e}")

            else:
                if os.path.exists(self.current_filename):
                    try:
                        current_size = os.path.getsize(self.current_filename)
                        img_size = os.path.getsize(img_path)
                        if self.max_size_bytes > 0 and (current_size + img_size + 10240) > self.max_size_bytes:
                            self._rotate_file()
                    except OSError:
                        pass

                if self.doc is None:
                    self.initialize_document()

                text_parts = []
                if self.log_window_titles and window_title:
                    text_parts.append(window_title)
                if self.append_sequence_number:
                    text_parts.append(str(count))

                if text_parts:
                    self.doc.add_paragraph(" ".join(text_parts))

                self.doc.add_picture(img_path, width=Inches(6))
                self.doc.add_paragraph("-" * 50)

                try:
                    self.doc.save(self.current_filename)
                    self.last_known_size_str = self.get_formatted_size(self.current_filename)
                    if self.gui_callback:
                        self.gui_callback("UPDATE_SIZE", self.last_known_size_str)
                    self.file_locked_warning_shown = False
                except PermissionError:
                    if not self.file_locked_warning_shown:
                        if self.gui_callback:
                            self.gui_callback("WARNING", "File Open in Word!", "Close to save")
                        self.file_locked_warning_shown = True

        except Exception as e:
            print(f"Save Error: {e}")

    def _process_undo(self):
        if self.session_count <= 0: return
        try:
            if self.image_paths:
                last_img = self.image_paths.pop()
                if os.path.exists(last_img):
                    try:
                        os.remove(last_img)
                    except OSError:
                        pass

            if self.save_mode == "folder":
                target_file = os.path.join(self.current_filename, f"{self.base_name_no_ext}_{self.session_count}.jpg")
                if os.path.exists(target_file):
                    try:
                        os.remove(target_file)
                    except OSError:
                        pass
                self.last_known_size_str = self.get_folder_size_str(self.current_filename)
                if self.gui_callback:
                    self.gui_callback("UNDO", self.session_count, self.last_known_size_str)

            else:
                if self.doc and len(self.doc.paragraphs) >= 2:
                    try:
                        p = self.doc.paragraphs[-1]
                        p._element.getparent().remove(p._element)
                        if self.doc.paragraphs:
                            p = self.doc.paragraphs[-1]
                            p._element.getparent().remove(p._element)
                        if self.doc.paragraphs:
                            last_p = self.doc.paragraphs[-1]
                            if last_p.text != "-" * 50 and not last_p.text.startswith("Screenshot Log"):
                                last_p._element.getparent().remove(last_p._element)
                    except Exception:
                        pass

                    try:
                        self.doc.save(self.current_filename)
                        self.last_known_size_str = self.get_formatted_size(self.current_filename)
                        self.file_locked_warning_shown = False
                    except PermissionError:
                        if self.gui_callback:
                            self.gui_callback("WARNING", "File Open in Word!", "Undo in memory only")

            self.session_count -= 1
            if self.gui_callback:
                self.gui_callback("UNDO", self.session_count, self.last_known_size_str)

        except Exception as e:
            print(f"Undo Error: {e}")

    def _rotate_file(self):
        if self.save_mode == "folder": return

        print(f"\n[!] Limit reached. Rotating file...")
        
        dir_name = os.path.dirname(self.current_filename)
        
        if not self.is_split_mode:
            self.is_split_mode = True
            part1_name = f"{self.base_name_no_ext}_Part1.docx"
            part1_path = os.path.join(dir_name, part1_name)
            
            part2_name = f"{self.base_name_no_ext}_Part2.docx"
            part2_path = os.path.join(dir_name, part2_name)
            
            self.doc = None 
            
            try:
                if not os.path.exists(part1_path):
                    os.rename(self.current_filename, part1_path)
            except OSError as e:
                print(f"Rename Error: {e}")
                
            self.current_filename = part2_path
            self.current_part = 2
        else:
            self.current_part += 1
            new_name = f"{self.base_name_no_ext}_Part{self.current_part}.docx"
            self.current_filename = os.path.join(dir_name, new_name)
        
        print(f"[i] Now saving to: {self.current_filename}")
        
        self.doc = Document()
        self.doc.add_heading(f'Screenshot Log - Part {self.current_part}', 0)
        try:
            self.doc.save(self.current_filename)
        except PermissionError:
            pass
        self.last_known_size_str = self.get_formatted_size(self.current_filename)
        
        if self.gui_callback:
            self.gui_callback("UPDATE_FILENAME", self.current_filename)

    def force_rotate(self):
        if not self.running or not self.current_filename: return False
        if self.save_mode == "folder": return False
        self._rotate_file()
        return True

    def cleanup_temp(self):
        if self.temp_dir and os.path.exists(self.temp_dir):
            try:
                shutil.rmtree(self.temp_dir)
                self.temp_dir = None
                self.image_paths = []
            except Exception as e:
                print(f"Error cleaning temp dir: {e}")

    def get_cleaned_window_title(self):
        try:
            hwnd = user32.GetForegroundWindow()
            length = user32.GetWindowTextLengthW(hwnd)
            buff = ctypes.create_unicode_buffer(length + 1)
            user32.GetWindowTextW(hwnd, buff, length + 1)
            full_title = buff.value

            clean_title = full_title.replace(" - Google Chrome", "") \
                .replace(" - Microsoft Edge", "") \
                .replace(" - Mozilla Firefox", "")
            return clean_title
        except Exception:
            return "Unknown Window"

    def get_unique_filename(self, base_name):
        if self.save_mode == "folder":
            full_path = os.path.join(self.save_directory, base_name)
            if not os.path.exists(full_path):
                return full_path
            counter = 1
            while True:
                new_name = f"{base_name}_{counter}"
                new_path = os.path.join(self.save_directory, new_name)
                if not os.path.exists(new_path):
                    return new_path
                counter += 1
        else:
            full_path = os.path.join(self.save_directory, base_name)
            if not full_path.lower().endswith(".docx"):
                full_path += ".docx"
            if not os.path.exists(full_path):
                return full_path
            name, ext = os.path.splitext(full_path)
            counter = 1
            while True:
                new_name = f"{name}_{counter}{ext}"
                if not os.path.exists(new_name):
                    return new_name
                counter += 1

    def get_formatted_size(self, filepath):
        if not os.path.exists(filepath): return "0 KB"
        size_bytes = os.path.getsize(filepath)
        if size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.2f} KB"
        return f"{size_bytes / (1024 * 1024):.2f} MB"

    def get_folder_size_str(self, folder_path):
        total_size = 0
        for dirpath, dirnames, filenames in os.walk(folder_path):
            for f in filenames:
                fp = os.path.join(dirpath, f)
                total_size += os.path.getsize(fp)
        return self._format_bytes(total_size)

    def _format_bytes(self, size_bytes):
        if size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.2f} KB"
        return f"{size_bytes / (1024 * 1024):.2f} MB"

    def initialize_document(self):
        if self.save_mode == "folder":
            if not os.path.exists(self.current_filename):
                os.makedirs(self.current_filename)
            self.last_known_size_str = self.get_folder_size_str(self.current_filename)
            return

        if os.path.exists(self.current_filename):
            try:
                self.doc = Document(self.current_filename)
            except Exception:
                self.doc = Document()
        else:
            self.doc = Document()
            if self.current_part > 1:
                self.doc.add_heading(f'Screenshot Log - Part {self.current_part}', 0)
            else:
                self.doc.add_heading('Screenshot Log', 0)
            try:
                self.doc.save(self.current_filename)
            except PermissionError:
                pass

        self.last_known_size_str = self.get_formatted_size(self.current_filename)

    def copy_dual_to_clipboard(self, img, filepaths):
        if not self.copy_files_option and not self.copy_image_option:
            return False

        try:
            h_dib = None
            h_drop = None

            if self.copy_image_option and img:
                output = io.BytesIO()
                img.convert("RGB").save(output, "BMP")
                dib_data = output.getvalue()[14:]
                output.close()
                dib_len = len(dib_data)

                h_dib = kernel32.GlobalAlloc(GPTR, dib_len)
                if h_dib:
                    ptr_dib = kernel32.GlobalLock(h_dib)
                    ctypes.memmove(ptr_dib, dib_data, dib_len)
                    kernel32.GlobalUnlock(h_dib)

            if self.copy_files_option and filepaths:
                if isinstance(filepaths, str):
                    filepaths = [filepaths]

                files_text = "\0".join([os.path.abspath(p) for p in filepaths]) + "\0\0"
                files_data = files_text.encode('utf-16le')

                pDropFiles = DROPFILES()
                pDropFiles.pFiles = ctypes.sizeof(DROPFILES)
                pDropFiles.pt = POINT(0, 0)
                pDropFiles.fNC = False
                pDropFiles.fWide = True

                drop_len = ctypes.sizeof(DROPFILES) + len(files_data)
                h_drop = kernel32.GlobalAlloc(GPTR, drop_len)
                if h_drop:
                    ptr_drop = kernel32.GlobalLock(h_drop)
                    ptr_addr = ptr_drop if isinstance(ptr_drop, int) else ptr_drop.value if ptr_drop else 0

                    if ptr_addr:
                        ctypes.memmove(ptr_drop, ctypes.byref(pDropFiles), ctypes.sizeof(DROPFILES))
                        ctypes.memmove(ptr_addr + ctypes.sizeof(DROPFILES), files_data, len(files_data))
                        kernel32.GlobalUnlock(h_drop)
                    else:
                        kernel32.GlobalFree(h_drop)
                        h_drop = None

            success = False
            for _ in range(10):
                if user32.OpenClipboard(None):
                    try:
                        user32.EmptyClipboard()
                        if h_dib:
                            user32.SetClipboardData(8, h_dib)
                        if h_drop:
                            user32.SetClipboardData(15, h_drop)
                        success = True
                    finally:
                        user32.CloseClipboard()
                    break
                time.sleep(0.01)

            if not success:
                if h_dib: kernel32.GlobalFree(h_dib)
                if h_drop: kernel32.GlobalFree(h_drop)

            return success
        except Exception as e:
            print(f"Dual Copy Error: {e}")
            return False

    def manual_copy_all(self):
        if not self.image_paths: return False

        total = len(self.image_paths)
        if self.gui_callback:
            self.gui_callback("PROGRESS_START", total)

        try:
            if self.copy_image_option:
                for i, path in enumerate(self.image_paths[:-1]):
                    if os.path.exists(path):
                        with Image.open(path) as img:
                            img.load()
                            self.copy_dual_to_clipboard(img.copy(), [path])
                            if self.gui_callback:
                                self.gui_callback("PROGRESS_UPDATE", i + 1)
                            time.sleep(1.0)

            last_img_path = self.image_paths[-1]
            if os.path.exists(last_img_path):
                with Image.open(last_img_path) as img:
                    img.load()
                    result = self.copy_dual_to_clipboard(img.copy(), self.image_paths)
                    if self.gui_callback:
                        self.gui_callback("PROGRESS_UPDATE", total)
                        self.gui_callback("PROGRESS_DONE", result, total)
                    return result

        except Exception as e:
            print(f"Manual Copy Error: {e}")
            if self.gui_callback:
                self.gui_callback("PROGRESS_DONE", False, 0)
        return False

    def take_screenshot(self):
        if not self.current_filename or not self.running: return

        try:
            img = ImageGrab.grab()
        except Exception as e:
            print(f"Capture Error: {e}")
            return

        self.session_count += 1
        count = self.session_count

        if self.gui_callback:
            self.gui_callback("NOTIFY", count, self.last_known_size_str)

        if self.auto_copy_clipboard:
            if not self.temp_dir or not os.path.exists(self.temp_dir):
                self.temp_dir = tempfile.mkdtemp(prefix="ScreenshotTool_")

            if self.save_mode == "folder":
                img_filename = f"{self.base_name_no_ext}_{count}.jpg"
            else:
                img_filename = f"screenshot_{count}.jpg"
                
            img_path = os.path.join(self.temp_dir, img_filename)

            threading.Thread(target=self._clipboard_worker, args=(img, img_path), daemon=True).start()

        window_title = self.get_cleaned_window_title() if self.log_window_titles else None
        self.save_queue.put((img, count, window_title))

    def _clipboard_worker(self, img, img_path):
        try:
            img.convert("RGB").save(img_path, "JPEG", quality=90)
            self.copy_dual_to_clipboard(img, [img_path])
        except Exception as e:
            print(f"Clipboard Worker Error: {e}")

    def undo_last_screenshot(self):
        if not self.current_filename or not self.running: return
        self.save_queue.put(("UNDO", None, None))


class ToonConfig:
    @staticmethod
    def load(filepath):
        config = {}
        if not os.path.exists(filepath): return config
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith('#'): continue
                    if ':' in line:
                        key, value = line.split(':', 1)
                        key, value = key.strip(), value.strip()
                        if value.lower() == 'true':
                            value = True
                        elif value.lower() == 'false':
                            value = False
                        config[key] = value
        except Exception as e:
            print(f"Error loading TOON config: {e}")
        return config

    @staticmethod
    def save(filepath, data):
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                for key, value in data.items():
                    f.write(f"{key}: {value}\n")
        except Exception as e:
            print(f"Error saving TOON config: {e}")


class MainWindow:
    def __init__(self, session):
        self.session = session
        self.config_file = "config.toon"
        self.gui_queue = queue.Queue()

        self.colors = {
            "bg": "#F5F5F5", "fg": "#000000", "accent": "#333333",
            "input_bg": "#FFFFFF", "input_fg": "#000000",
            "btn_bg": "#E0E0E0", "btn_fg": "#000000",
            "btn_hover": "#D6D6D6", "btn_active": "#BDBDBD",
            "success": "#2E7D32", "error": "#C62828",
            "warning": "#F57F17"
        }

        ctypes.windll.kernel32.GlobalLock.argtypes = [wintypes.HGLOBAL]
        ctypes.windll.kernel32.GlobalLock.restype = ctypes.c_void_p

        self.root = tk.Tk()
        self.root.title("Screenshot Tool")
        self.root.geometry("600x700")
        self.root.configure(bg=self.colors["bg"])

        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width / 2) - (600 / 2)
        y = (screen_height / 2) - (700 / 2)
        self.root.geometry(f"600x700+{int(x)}+{int(y)}")

        self.notif_window = tk.Toplevel(self.root)
        self.notif_window.withdraw()
        self.notif_window.overrideredirect(True)
        self.notif_window.attributes("-topmost", True)
        self.notif_window.configure(bg="white")

        self.notif_frame = tk.Frame(self.notif_window, bg="white",
                                    highlightbackground="#2E7D32", highlightthickness=1)
        self.notif_frame.pack(fill=tk.BOTH, expand=True)

        self.lbl_notif_title = tk.Label(self.notif_frame, text="",
                                        fg="#2E7D32", bg="white",
                                        font=("Segoe UI", 10, "bold"))
        self.lbl_notif_title.pack(pady=(2, 0))

        self.lbl_notif_info = tk.Label(self.notif_frame, text="", fg="#2E7D32",
                                       bg="white", font=("Segoe UI", 8))
        self.lbl_notif_info.pack(pady=(0, 2))

        self.hide_timer = None

        main_frame = tk.Frame(self.root, bg=self.colors["bg"], padx=30, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        tk.Label(main_frame, text="Screenshot Tool", font=("Segoe UI", 18, "bold"),
                 bg=self.colors["bg"], fg=self.colors["accent"]).pack(pady=(0, 15))

        self.create_label(main_frame, "Document Name")
        self.entry_name = self.create_entry(main_frame)
        self.entry_name.insert(0, "screenshots")

        self.create_label(main_frame, "Max File Size (MB) - 0 for Unlimited")
        self.entry_size = self.create_entry(main_frame)
        self.entry_size.insert(0, "0")

        self.create_label(main_frame, "Save Directory")
        dir_frame = tk.Frame(main_frame, bg=self.colors["bg"])
        dir_frame.pack(fill=tk.X, pady=(0, 10))

        self.entry_dir = self.create_entry(dir_frame, pack=False)
        self.entry_dir.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))

        self.btn_browse = self.create_button(dir_frame, "Browse", self.browse_directory, width=10)
        self.btn_browse.pack(side=tk.RIGHT)

        self.create_label(main_frame, "Save Mode")
        self.combo_mode = ttk.Combobox(main_frame, values=["Word Document (.docx)", "Raw Images (Folder)"], state="readonly")
        self.combo_mode.current(0)
        self.combo_mode.pack(fill=tk.X, pady=(0, 10))

        self.var_log_title = tk.BooleanVar()
        self.chk_title = self.create_checkbox(main_frame, "Log Active Window Title", self.var_log_title)

        self.var_append_num = tk.BooleanVar(value=True)
        self.chk_num = self.create_checkbox(main_frame, "Append Screenshot Number", self.var_append_num)

        self.var_auto_copy = tk.BooleanVar(value=False)
        self.chk_auto_copy = self.create_checkbox(main_frame, "Auto-copy to Clipboard", self.var_auto_copy)

        self.var_copy_files = tk.BooleanVar(value=True)
        frame_cb1 = tk.Frame(main_frame, bg=self.colors["bg"])
        frame_cb1.pack(fill=tk.X, pady=1)
        self.chk_copy_files = self.create_checkbox(frame_cb1, "Include Files", self.var_copy_files, pack=False)
        self.chk_copy_files.pack(side=tk.LEFT)
        tk.Label(frame_cb1, text="(Paste in Explorer/Email)", fg="gray", bg=self.colors["bg"],
                 font=("Segoe UI", 8)).pack(side=tk.LEFT, padx=5)

        self.var_copy_image = tk.BooleanVar(value=True)
        frame_cb2 = tk.Frame(main_frame, bg=self.colors["bg"])
        frame_cb2.pack(fill=tk.X, pady=1)
        self.chk_copy_image = self.create_checkbox(frame_cb2, "Include Image", self.var_copy_image, pack=False)
        self.chk_copy_image.pack(side=tk.LEFT)
        tk.Label(frame_cb2, text="(Visible in Win+V History)", fg="gray", bg=self.colors["bg"],
                 font=("Segoe UI", 8)).pack(side=tk.LEFT, padx=5)

        # Hotkey Info
        info_text = "~ (Tilde): Take Screenshot | Ctrl+Alt+~: Undo Last"
        tk.Label(main_frame, text=info_text, fg="gray", bg=self.colors["bg"], font=("Segoe UI", 8, "italic")).pack(pady=(10, 5))

        info_frame = tk.Frame(main_frame, bg=self.colors["bg"])
        info_frame.pack(fill=tk.X, pady=(5, 5))

        self.status_label = tk.Label(info_frame, text="Ready to Start", bg=self.colors["bg"],
                                     fg=self.colors["fg"], font=("Segoe UI", 9))
        self.status_label.pack(side=tk.LEFT)

        self.count_label = tk.Label(info_frame, text="Count: 0", bg=self.colors["bg"],
                                    fg=self.colors["accent"], font=("Segoe UI", 9, "bold"))
        self.count_label.pack(side=tk.RIGHT, padx=(10, 0))

        self.size_label = tk.Label(info_frame, text="Size: 0 KB", bg=self.colors["bg"],
                                   fg=self.colors["accent"], font=("Segoe UI", 9, "bold"))
        self.size_label.pack(side=tk.RIGHT)

        self.btn_start = self.create_button(main_frame, "START SESSION", self.toggle_session,
                                            bg=self.colors["accent"], fg="#FFFFFF", height=2)
        self.btn_start.pack(fill=tk.X, pady=(5, 5))
        self.btn_start.configure(font=("Segoe UI", 10, "bold"))

        action_frame = tk.Frame(main_frame, bg=self.colors["bg"])
        action_frame.pack(fill=tk.X, pady=5)

        self.btn_rotate = self.create_button(action_frame, "SPLIT FILE", self.manual_rotate,
                                             bg=self.colors["btn_bg"], fg=self.colors["btn_fg"], height=2)
        self.btn_rotate.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
        self.btn_rotate.configure(font=("Segoe UI", 9, "bold"))

        self.btn_copy = self.create_button(action_frame, "COPY ALL IMAGES", self.copy_to_clipboard,
                                           bg=self.colors["btn_bg"], fg=self.colors["btn_fg"], height=2)
        self.btn_copy.pack(side=tk.RIGHT, fill=tk.X, expand=True, padx=(5, 0))
        self.btn_copy.configure(font=("Segoe UI", 9, "bold"))

        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(main_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X, pady=(10, 0))
        self.progress_bar.pack_forget()

        self.load_settings()

        self.session.set_callback(self.queue_notification)
        self.check_gui_queue()
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def create_label(self, parent, text):
        lbl = tk.Label(parent, text=text, bg=self.colors["bg"], fg=self.colors["fg"], font=("Segoe UI", 9))
        lbl.pack(anchor="w", pady=(0, 2))
        return lbl

    def create_entry(self, parent, pack=True):
        entry = tk.Entry(parent, bg=self.colors["input_bg"], fg=self.colors["input_fg"],
                         insertbackground=self.colors["fg"], relief="flat", font=("Segoe UI", 9))
        if pack: entry.pack(fill=tk.X, ipady=4, pady=(0, 10))
        return entry

    def create_button(self, parent, text, command, bg=None, fg=None, width=None, height=1):
        if bg is None: bg = self.colors["btn_bg"]
        if fg is None: fg = self.colors["btn_fg"]
        btn = tk.Button(parent, text=text, command=command, bg=bg, fg=fg,
                        activebackground=self.colors["btn_hover"], activeforeground=fg,
                        relief="flat", borderwidth=0, font=("Segoe UI", 9),
                        width=width, height=height, cursor="hand2")
        return btn

    def create_checkbox(self, parent, text, variable, pack=True):
        chk = tk.Checkbutton(parent, text=text, variable=variable,
                             bg=self.colors["bg"], fg=self.colors["fg"],
                             selectcolor=self.colors["input_bg"],
                             activebackground=self.colors["bg"], activeforeground=self.colors["fg"],
                             font=("Segoe UI", 9))
        if pack:
            chk.pack(anchor="w", pady=1)
        return chk

    def browse_directory(self):
        directory = filedialog.askdirectory()
        if directory:
            self.entry_dir.delete(0, tk.END)
            self.entry_dir.insert(0, directory)

    def load_settings(self):
        default_dir = os.path.join(os.path.expanduser("~"), "Desktop", "Evidence")
        data = ToonConfig.load(self.config_file)
        self.entry_name.delete(0, tk.END)
        self.entry_name.insert(0, data.get("filename", "screenshots"))
        self.entry_size.delete(0, tk.END)
        self.entry_size.insert(0, str(data.get("max_size", "0")))
        self.entry_dir.delete(0, tk.END)
        self.entry_dir.insert(0, data.get("save_dir", default_dir))
        self.var_log_title.set(data.get("log_title", False))
        self.var_append_num.set(data.get("append_num", True))
        self.var_auto_copy.set(data.get("auto_copy", False))
        self.var_copy_files.set(data.get("copy_files", True))
        self.var_copy_image.set(data.get("copy_image", True))
        
        mode = data.get("save_mode", "docx")
        if mode == "folder":
            self.combo_mode.current(1)
        else:
            self.combo_mode.current(0)

    def save_settings(self):
        mode = "docx" if self.combo_mode.current() == 0 else "folder"
        data = {
            "filename": self.entry_name.get(),
            "max_size": self.entry_size.get(),
            "save_dir": self.entry_dir.get(),
            "log_title": self.var_log_title.get(),
            "append_num": self.var_append_num.get(),
            "auto_copy": self.var_auto_copy.get(),
            "copy_files": self.var_copy_files.get(),
            "copy_image": self.var_copy_image.get(),
            "save_mode": mode
        }
        ToonConfig.save(self.config_file, data)

    def queue_notification(self, *args):
        self.gui_queue.put(args)

    def check_gui_queue(self):
        try:
            while True:
                args = self.gui_queue.get_nowait()
                action = args[0]

                if action == "NOTIFY":
                    count, size_str = args[1], args[2]
                    self.show_notification(f"✓ #{count}", size_str)
                    self.count_label.config(text=f"Count: {count}")

                elif action == "UPDATE_SIZE":
                    size_str = args[1]
                    self.size_label.config(text=f"Size: {size_str}")
                    if self.notif_window.state() == "normal":
                        self.lbl_notif_info.config(text=f"{size_str}")
                
                elif action == "UPDATE_FILENAME":
                    new_path = args[1]
                    date_str = datetime.datetime.now().strftime("%d-%m-%Y")
                    display_path = f".../{date_str}/{os.path.basename(new_path)}"
                    self.status_label.config(text=f"Active | {display_path}")

                elif action == "UNDO":
                    count, size_str = args[1], args[2]
                    self.show_notification(f"↺ Undone", size_str)
                    self.size_label.config(text=f"Size: {size_str}")
                    self.count_label.config(text=f"Count: {count}")

                elif action == "WARNING":
                    title, msg = args[1], args[2]
                    self.show_notification(f"⚠ {title}", msg)
                    self.status_label.config(text=f"⚠ {title}", fg=self.colors["error"])

                elif action == "PROGRESS_START":
                    total = args[1]
                    self.progress_bar.config(maximum=total, value=0)
                    self.progress_bar.pack(fill=tk.X, pady=(10, 0))
                    self.status_label.config(text="Copying to Clipboard...", fg=self.colors["accent"])
                    self.btn_copy.config(state='disabled')

                elif action == "PROGRESS_UPDATE":
                    val = args[1]
                    self.progress_var.set(val)

                elif action == "PROGRESS_DONE":
                    success = args[1]
                    total = args[2] if len(args) > 2 else 0
                    self.progress_bar.pack_forget()
                    self.btn_copy.config(state='normal')
                    if success:
                        self.status_label.config(text=f"Copied {total} images!", fg=self.colors["success"])
                    else:
                        self.status_label.config(text="Copy Failed", fg=self.colors["error"])

        except queue.Empty:
            pass
        self.root.after(20, self.check_gui_queue)

    def show_notification(self, title, size_str):
        self.lbl_notif_title.config(text=title)
        self.lbl_notif_info.config(text=f"{size_str}")
        self.size_label.config(text=f"Size: {size_str}")

        width, height = 150, 45
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        self.notif_window.geometry(f"{width}x{height}+{screen_width - width - 20}+{screen_height - height - 60}")

        self.notif_window.deiconify()

        if self.hide_timer:
            self.root.after_cancel(self.hide_timer)
        self.hide_timer = self.root.after(1500, self.notif_window.withdraw)

    def toggle_session(self):
        if not self.session.running:
            self.start_app()
        else:
            self.stop_app()

    def start_app(self):
        self.save_settings()
        self.session.cleanup_temp()

        self.session.session_count = 0
        self.session.current_part = 1
        self.session.is_split_mode = False
        self.session.doc = None
        self.session.file_locked_warning_shown = False
        self.session.last_known_size_str = "0 KB"

        # Determine Mode
        mode_idx = self.combo_mode.current()
        self.session.save_mode = "docx" if mode_idx == 0 else "folder"

        base_dir = self.entry_dir.get().strip()
        if not base_dir: base_dir = os.path.join(os.path.expanduser("~"), "Desktop", "Evidence")
        date_str = datetime.datetime.now().strftime("%d-%m-%Y")
        save_dir = os.path.join(base_dir, date_str)

        if not os.path.exists(save_dir):
            try:
                os.makedirs(save_dir)
            except OSError as e:
                self.status_label.config(text=f"Error: {e}", fg=self.colors["error"])
                return

        self.session.save_directory = save_dir
        self.session.temp_dir = tempfile.mkdtemp(prefix="ScreenshotTool_")

        user_input = self.entry_name.get().strip()
        if not user_input: user_input = "screenshots"
        self.session.current_filename = self.session.get_unique_filename(user_input)
        
        # FIX: Use basename to avoid full path in filename construction
        self.session.base_name_no_ext = os.path.splitext(os.path.basename(self.session.current_filename))[0]

        try:
            mb = float(self.entry_size.get().strip())
            self.session.max_size_bytes = int(mb * 1024 * 1024)
        except ValueError:
            self.session.max_size_bytes = 0

        self.session.log_window_titles = self.var_log_title.get()
        self.session.append_sequence_number = self.var_append_num.get()
        self.session.auto_copy_clipboard = self.var_auto_copy.get()
        self.session.copy_files_option = self.var_copy_files.get()
        self.session.copy_image_option = self.var_copy_image.get()
        self.session.running = True

        self.session.start_worker()
        self.session.initialize_document()

        keyboard.add_hotkey('~', self.session.take_screenshot, suppress=True)
        keyboard.add_hotkey('ctrl+alt+~', self.undo_action, suppress=True)

        self.btn_start.config(text="STOP SESSION", bg=self.colors["error"], fg="#FFFFFF")
        self.btn_rotate.config(state='normal')
        self.entry_name.config(state='disabled')
        self.entry_size.config(state='disabled')
        self.entry_dir.config(state='disabled')
        self.btn_browse.config(state='disabled')
        self.combo_mode.config(state='disabled')
        self.chk_title.config(state='disabled')
        self.chk_num.config(state='disabled')
        self.chk_auto_copy.config(state='disabled')
        self.chk_copy_files.config(state='disabled')
        self.chk_copy_image.config(state='disabled')

        display_path = f".../{date_str}/{os.path.basename(self.session.current_filename)}"
        self.status_label.config(text=f"Active | {display_path}", fg=self.colors["success"])
        self.size_label.config(text="Size: 0 KB")
        self.count_label.config(text="Count: 0")

        print("\n" + "=" * 40)
        print("      SESSION STARTED")
        print("=" * 40)
        print(f"[OK] File: {self.session.current_filename}")
        print(f"[OK] Press '~' (tilde) to take a screenshot.")

    def stop_app(self):
        self.session.running = False
        keyboard.unhook_all_hotkeys()

        def wait_for_queue():
            if not self.session.save_queue.empty():
                self.status_label.config(text="Saving pending screenshots...", fg=self.colors["warning"])
                self.root.after(100, wait_for_queue)
            else:
                self.finish_stop()

        wait_for_queue()

    def finish_stop(self):
        self.btn_start.config(text="START SESSION", bg=self.colors["accent"], fg="#FFFFFF")
        self.btn_rotate.config(state='disabled')
        self.entry_name.config(state='normal')
        self.entry_size.config(state='normal')
        self.entry_dir.config(state='normal')
        self.btn_browse.config(state='normal')
        self.combo_mode.config(state='readonly')
        self.chk_title.config(state='normal')
        self.chk_num.config(state='normal')
        self.chk_auto_copy.config(state='normal')
        self.chk_copy_files.config(state='normal')
        self.chk_copy_image.config(state='normal')
        self.status_label.config(text="Session Stopped", fg=self.colors["warning"])
        print("\n[OK] Session Stopped.")

    def manual_rotate(self):
        if self.session.force_rotate():
            self.status_label.config(text="File Split Manually", fg=self.colors["success"])
            date_str = datetime.datetime.now().strftime("%d-%m-%Y")
            display_path = f".../{date_str}/{os.path.basename(self.session.current_filename)}"
            self.status_label.config(text=f"Active | {display_path}")
            self.size_label.config(text="Size: 0 KB")
            self.count_label.config(text=f"Count: {self.session.session_count}")

    def undo_action(self):
        self.session.undo_last_screenshot()

    def copy_to_clipboard(self):
        if not self.session.image_paths:
            self.status_label.config(text="No images to copy!", fg=self.colors["warning"])
            return

        threading.Thread(target=self.session.manual_copy_all, daemon=True).start()

    def on_close(self):
        self.save_settings()
        if self.session.running:
            self.stop_app()
        self.session.cleanup_temp()
        self.root.destroy()


def main():
    session = ScreenshotSession()
    app = MainWindow(session)
    try:
        app.root.mainloop()
    except KeyboardInterrupt:
        print("\nProgram interrupted by user.")
        if session.running:
            session.running = False
            session.cleanup_temp()


if __name__ == "__main__":
    main()