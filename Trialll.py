from PIL import Image, ImageGrab
from docx import Document
from docx.shared import Inches
import os
import datetime
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
import threading
import ctypes
from ctypes import wintypes
import tempfile
import shutil
import queue
import time
import io

# --- CTYPES SETUP (Windows API) ---
kernel32 = ctypes.windll.kernel32
user32 = ctypes.windll.user32

SIZE_T = ctypes.c_size_t
HGLOBAL = wintypes.HGLOBAL
LPVOID = ctypes.c_void_p
BOOL = wintypes.BOOL
UINT = wintypes.UINT
HANDLE = wintypes.HANDLE
HWND = wintypes.HWND
DWORD = wintypes.DWORD

GMEM_FIXED = 0x0000
GMEM_ZEROINIT = 0x0040
GPTR = GMEM_FIXED | GMEM_ZEROINIT

# Clipboard Functions
kernel32.GlobalAlloc.argtypes = [UINT, SIZE_T]
kernel32.GlobalAlloc.restype = HGLOBAL
kernel32.GlobalLock.argtypes = [HGLOBAL]
kernel32.GlobalLock.restype = LPVOID
kernel32.GlobalUnlock.argtypes = [HGLOBAL]
kernel32.GlobalUnlock.restype = BOOL
kernel32.GlobalFree.argtypes = [HGLOBAL]
kernel32.GlobalFree.restype = HGLOBAL
kernel32.GetCurrentThreadId.restype = DWORD

user32.OpenClipboard.argtypes = [HWND]
user32.OpenClipboard.restype = BOOL
user32.EmptyClipboard.argtypes = []
user32.EmptyClipboard.restype = BOOL
user32.SetClipboardData.argtypes = [UINT, HANDLE]
user32.SetClipboardData.restype = HANDLE
user32.CloseClipboard.argtypes = []
user32.CloseClipboard.restype = BOOL

# Hotkey Functions
user32.RegisterHotKey.argtypes = [HWND, ctypes.c_int, UINT, UINT]
user32.RegisterHotKey.restype = BOOL
user32.UnregisterHotKey.argtypes = [HWND, ctypes.c_int]
user32.UnregisterHotKey.restype = BOOL
user32.GetMessageW.argtypes = [ctypes.POINTER(wintypes.MSG), HWND, UINT, UINT]
user32.GetMessageW.restype = BOOL
user32.TranslateMessage.argtypes = [ctypes.POINTER(wintypes.MSG)]
user32.DispatchMessageW.argtypes = [ctypes.POINTER(wintypes.MSG)]
user32.PostThreadMessageW.argtypes = [DWORD, UINT, wintypes.WPARAM, wintypes.LPARAM]
user32.PostThreadMessageW.restype = BOOL

# Constants
MOD_ALT = 0x0001
MOD_CONTROL = 0x0002
VK_OEM_3 = 0xC0  # The '~' key
WM_HOTKEY = 0x0312
WM_USER = 0x0400
WM_STOP_LISTENER = WM_USER + 1


class POINT(ctypes.Structure):
    _fields_ = [("x", ctypes.c_long), ("y", ctypes.c_long)]


class DROPFILES(ctypes.Structure):
    _fields_ = [
        ("pFiles", wintypes.DWORD),
        ("pt", POINT),
        ("fNC", wintypes.BOOL),
        ("fWide", wintypes.BOOL),
    ]


# --- LOGIC CLASSES ---

class HotkeyListener:
    def __init__(self, callback_capture, callback_undo):
        self.callback_capture = callback_capture
        self.callback_undo = callback_undo
        self.thread = None
        self.thread_id = None
        self.running = False

    def start(self):
        if self.thread is None or not self.thread.is_alive():
            self.running = True
            self.thread = threading.Thread(target=self._loop, daemon=True)
            self.thread.start()

    def stop(self):
        self.running = False
        if self.thread_id:
            user32.PostThreadMessageW(self.thread_id, WM_STOP_LISTENER, 0, 0)

    def _loop(self):
        self.thread_id = kernel32.GetCurrentThreadId()
        HOTKEY_CAPTURE = 1
        HOTKEY_UNDO = 2

        user32.RegisterHotKey(None, HOTKEY_CAPTURE, 0x0000, VK_OEM_3)
        user32.RegisterHotKey(None, HOTKEY_UNDO, MOD_CONTROL | MOD_ALT, VK_OEM_3)

        msg = wintypes.MSG()
        while user32.GetMessageW(ctypes.byref(msg), None, 0, 0) != 0:
            if msg.message == WM_HOTKEY:
                if msg.wParam == HOTKEY_CAPTURE:
                    if self.callback_capture: self.callback_capture()
                elif msg.wParam == HOTKEY_UNDO:
                    if self.callback_undo: self.callback_undo()
            elif msg.message == WM_STOP_LISTENER:
                break

            user32.TranslateMessage(ctypes.byref(msg))
            user32.DispatchMessageW(ctypes.byref(msg))

        user32.UnregisterHotKey(None, HOTKEY_CAPTURE)
        user32.UnregisterHotKey(None, HOTKEY_UNDO)


class ScreenshotSession:
    def __init__(self, config_data, callback_queue):
        self.config = config_data
        self.gui_queue = callback_queue

        self.base_name_no_ext = ""
        self.current_filename = ""
        self.session_count = 0
        self.max_size_bytes = 0
        self.current_part = 1
        self.is_split_mode = False
        self.image_paths = []
        self.temp_dir = None
        self.doc = None
        self.save_queue = queue.Queue()
        self.worker_thread = None
        self.file_locked_warning_shown = False
        self.last_known_size_str = "0 KB"
        self.running = True
        self.status = "Active"

        self._initialize()

    def _initialize(self):
        base_dir = self.config['save_dir']
        filename_input = self.config['filename']

        if not os.path.exists(base_dir):
            try:
                os.makedirs(base_dir)
            except:
                pass

        full_path = os.path.join(base_dir, filename_input)

        if self.config['save_mode'] == 'folder':
            if not os.path.exists(full_path):
                self.current_filename = full_path
            else:
                c = 1
                while True:
                    new_path = f"{full_path}_{c}"
                    if not os.path.exists(new_path):
                        self.current_filename = new_path
                        break
                    c += 1
            self.base_name_no_ext = os.path.basename(self.current_filename)
        else:
            c = 0
            while True:
                name = filename_input if c == 0 else f"{filename_input}_{c}"
                f_path = os.path.join(base_dir, name + ".docx")
                if not os.path.exists(f_path):
                    self.current_filename = f_path
                    self.base_name_no_ext = name
                    break
                c += 1

        try:
            mb = float(self.config['max_size'])
            self.max_size_bytes = int(mb * 1024 * 1024)
        except:
            self.max_size_bytes = 0

        self.temp_dir = tempfile.mkdtemp(prefix="ShotTool_")

        if self.config['save_mode'] == 'folder':
            if not os.path.exists(self.current_filename): os.makedirs(self.current_filename)
        else:
            self.doc = Document()
            try:
                self.doc.save(self.current_filename)
            except:
                pass

        self.start_worker()

    def start_worker(self):
        if self.worker_thread is None or not self.worker_thread.is_alive():
            self.worker_thread = threading.Thread(target=self._worker_loop, daemon=True)
            self.worker_thread.start()

    def stop_worker(self):
        self.running = False

    def capture(self):
        if not self.running: return
        try:
            img = ImageGrab.grab()
        except:
            return

        self.session_count += 1
        window_title = self.get_cleaned_window_title() if self.config['log_title'] else None

        self.gui_queue.put(("NOTIFY", self.current_filename, self.session_count, self.last_known_size_str))

        if self.config['auto_copy']:
            path = os.path.join(self.temp_dir, f"clip_{self.session_count}.jpg")
            threading.Thread(target=self._clipboard_worker, args=(img, path), daemon=True).start()

        self.save_queue.put((img, self.session_count, window_title))

    def undo(self):
        if not self.running: return
        self.save_queue.put(("UNDO", None, None))

    def manual_rotate(self):
        if self.config['save_mode'] == "folder": return
        self._rotate_file()

    def _worker_loop(self):
        while True:
            try:
                try:
                    task = self.save_queue.get(timeout=0.1)
                except queue.Empty:
                    if not self.running and self.save_queue.empty(): break
                    continue

                if task[0] == "UNDO":
                    self._process_undo()
                else:
                    self._process_save(task[0], task[1], task[2])
                self.save_queue.task_done()
            except Exception as e:
                print(f"Worker Err: {e}")

    def _process_save(self, img, count, window_title):
        try:
            if self.config['save_mode'] == "folder":
                fname = f"{self.base_name_no_ext}_{count}.jpg"
            else:
                fname = f"screen_{count}.jpg"

            img_path = os.path.join(self.temp_dir, fname)
            img.save(img_path, "JPEG", quality=90)
            self.image_paths.append(img_path)

            if self.config['save_mode'] == "folder":
                target = os.path.join(self.current_filename, fname)
                shutil.copyfile(img_path, target)
                self.last_known_size_str = self.get_folder_size_str(self.current_filename)
            else:
                if os.path.exists(self.current_filename):
                    curr = os.path.getsize(self.current_filename)
                    if self.max_size_bytes > 0 and (curr + os.path.getsize(img_path)) > self.max_size_bytes:
                        self._rotate_file()

                if not self.doc: self.doc = Document(self.current_filename)

                txt = ""
                if window_title and self.config['append_num']:
                    txt = f"{window_title} {count}"
                elif window_title:
                    txt = window_title
                elif self.config['append_num']:
                    txt = str(count)

                if txt: self.doc.add_paragraph(txt)
                self.doc.add_picture(img_path, width=Inches(6))
                self.doc.add_paragraph("-" * 50)

                try:
                    self.doc.save(self.current_filename)
                    self.last_known_size_str = self.get_formatted_size(self.current_filename)
                    self.file_locked_warning_shown = False
                except PermissionError:
                    if not self.file_locked_warning_shown:
                        self.gui_queue.put(("WARNING", "File Locked", "Close Word to save"))
                        self.file_locked_warning_shown = True

            self.gui_queue.put(("UPDATE_SESSION", self.current_filename, self.session_count, self.last_known_size_str))

        except Exception as e:
            print(f"Save Err: {e}")

    def _process_undo(self):
        if self.session_count <= 0: return
        try:
            if self.image_paths:
                p = self.image_paths.pop()
                if os.path.exists(p): os.remove(p)

            if self.config['save_mode'] == 'folder':
                f = os.path.join(self.current_filename, f"{self.base_name_no_ext}_{self.session_count}.jpg")
                if os.path.exists(f): os.remove(f)
                self.last_known_size_str = self.get_folder_size_str(self.current_filename)
            else:
                if self.doc and len(self.doc.paragraphs) >= 2:
                    try:
                        p = self.doc.paragraphs[-1]
                        p._element.getparent().remove(p._element)
                        if self.doc.paragraphs:
                            p = self.doc.paragraphs[-1]
                            p._element.getparent().remove(p._element)
                        if self.doc.paragraphs:
                            last = self.doc.paragraphs[-1]
                            if last.text != "-" * 50: last._element.getparent().remove(last._element)
                        self.doc.save(self.current_filename)
                        self.last_known_size_str = self.get_formatted_size(self.current_filename)
                    except:
                        pass

            self.session_count -= 1
            self.gui_queue.put(("UNDO", self.current_filename, self.session_count, self.last_known_size_str))
        except:
            pass

    def cleanup(self, delete_files=False):
        self.stop_worker()
        if self.temp_dir and os.path.exists(self.temp_dir):
            try:
                shutil.rmtree(self.temp_dir)
            except:
                pass

        if delete_files and self.current_filename and os.path.exists(self.current_filename):
            try:
                if self.config['save_mode'] == 'folder':
                    shutil.rmtree(self.current_filename)
                else:
                    os.remove(self.current_filename)
            except Exception as e:
                print(f"Cleanup Error: {e}")

    def _rotate_file(self):
        dir_n = os.path.dirname(self.current_filename)
        base = self.base_name_no_ext.split("_Part")[0]
        c = 1
        while True:
            nn = f"{base}_Part{c}.docx"
            fp = os.path.join(dir_n, nn)
            if not os.path.exists(fp):
                try:
                    os.rename(self.current_filename, fp)
                except:
                    pass
                self.current_filename = os.path.join(dir_n, f"{base}_Part{c + 1}.docx")
                break
            c += 1
        self.doc = Document()
        self.doc.save(self.current_filename)
        self.last_known_size_str = "0 KB"
        self.gui_queue.put(("UPDATE_FILENAME", self.current_filename))

    def get_cleaned_window_title(self):
        try:
            hwnd = user32.GetForegroundWindow()
            length = user32.GetWindowTextLengthW(hwnd)
            buff = ctypes.create_unicode_buffer(length + 1)
            user32.GetWindowTextW(hwnd, buff, length + 1)
            full = buff.value
            return full.replace(" - Google Chrome", "").replace(" - Microsoft Edge", "")
        except:
            return "Unknown"

    def get_formatted_size(self, fp):
        if not os.path.exists(fp): return "0 KB"
        s = os.path.getsize(fp)
        return f"{s / 1024:.2f} KB" if s < 1048576 else f"{s / 1048576:.2f} MB"

    def get_folder_size_str(self, fp):
        t = 0
        for r, d, f in os.walk(fp):
            for file in f: t += os.path.getsize(os.path.join(r, file))
        return f"{t / 1024:.2f} KB" if t < 1048576 else f"{t / 1048576:.2f} MB"

    def _clipboard_worker(self, img, path):
        img.convert("RGB").save(path, "JPEG")
        self.copy_dual(img, [path])

    def copy_dual(self, img, paths):
        do_image = self.config.get('copy_image', True)
        do_files = self.config.get('copy_files', True)
        if not do_image and not do_files: return

        try:
            h_dib = None
            h_drop = None

            if do_image and img:
                output = io.BytesIO()
                img.convert("RGB").save(output, "BMP")
                data = output.getvalue()[14:]
                output.close()
                h_dib = kernel32.GlobalAlloc(GPTR, len(data))
                if h_dib:
                    p_dib = kernel32.GlobalLock(h_dib)
                    ctypes.memmove(p_dib, data, len(data))
                    kernel32.GlobalUnlock(h_dib)

            if do_files and paths:
                f_txt = "\0".join([os.path.abspath(p) for p in paths]) + "\0\0"
                f_dat = f_txt.encode('utf-16le')
                h_drop = kernel32.GlobalAlloc(GPTR, 20 + len(f_dat))
                if h_drop:
                    p_drop = kernel32.GlobalLock(h_drop)
                    header = b'\x14\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x01\x00\x00\x00'
                    ctypes.memmove(p_drop, header, 20)
                    ctypes.memmove(p_drop + 20, f_dat, len(f_dat))
                    kernel32.GlobalUnlock(h_drop)

            for _ in range(10):
                if user32.OpenClipboard(None):
                    try:
                        user32.EmptyClipboard()
                        if h_dib: user32.SetClipboardData(8, h_dib)
                        if h_drop: user32.SetClipboardData(15, h_drop)
                    finally:
                        user32.CloseClipboard()
                    break
                time.sleep(0.01)
        except:
            pass


class ToonConfig:
    @staticmethod
    def load(filepath):
        config = {}
        if not os.path.exists(filepath): return config
        try:
            with open(filepath, 'r') as f:
                for line in f:
                    if ':' in line:
                        k, v = line.split(':', 1)
                        config[k.strip()] = v.strip() == 'True' if v.strip() in ['True', 'False'] else v.strip()
        except:
            pass
        return config

    @staticmethod
    def save(filepath, data):
        try:
            with open(filepath, 'w') as f:
                for k, v in data.items(): f.write(f"{k}: {v}\n")
        except:
            pass


# --- UI CLASS ---
class ModernUI:
    def __init__(self):
        self.config_file = "config.toon"
        self.gui_queue = queue.Queue()
        self.sessions = {}
        self.active_session_key = None
        self.original_base_dir = None
        self.hide_timer = None

        # MODERN THEME COLORS (VS Code Dark Style)
        self.colors = {
            "bg": "#1e1e1e",
            "fg": "#d4d4d4",
            "sidebar": "#252526",
            "accent": "#007acc",
            "accent_hover": "#0062a3",
            "danger": "#f44336",
            "danger_hover": "#d32f2f",
            "success": "#4caf50",
            "success_hover": "#388e3c",
            "warning": "#FF9800",
            "entry_bg": "#3c3c3c",
            "entry_fg": "#ffffff",
            "border": "#3e3e42",
            "disabled": "#5a5a5a"
        }

        self.root = tk.Tk()
        self.root.title("Screenshot Tool Pro")
        self.root.configure(bg=self.colors['bg'])

        conf = ToonConfig.load(self.config_file)
        w, h = int(conf.get("w", 800)), int(conf.get("h", 600))
        self.root.geometry(f"{w}x{h}")

        # Setup Styles
        style = ttk.Style()
        style.theme_use('clam')

        # Treeview Style
        style.configure("Treeview",
                        background=self.colors['bg'],
                        foreground=self.colors['fg'],
                        fieldbackground=self.colors['bg'],
                        borderwidth=0)
        style.configure("Treeview.Heading",
                        background=self.colors['sidebar'],
                        foreground=self.colors['fg'],
                        relief="flat")
        style.map("Treeview", background=[('selected', self.colors['accent'])])

        # --- LAYOUT ---

        # 1. Sidebar (Session Manager)
        sidebar = tk.Frame(self.root, bg=self.colors['sidebar'], width=280)
        sidebar.pack(side=tk.LEFT, fill=tk.Y)
        sidebar.pack_propagate(False)

        tk.Label(sidebar, text="SESSIONS", bg=self.colors['sidebar'], fg=self.colors['fg'],
                 font=("Segoe UI", 10, "bold")).pack(pady=(20, 10), padx=10, anchor="w")

        # Treeview
        tree_frame = tk.Frame(sidebar, bg=self.colors['sidebar'])
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=10)

        self.tree = ttk.Treeview(tree_frame, columns=("status", "count"), show="tree", selectmode="browse")
        self.tree.column("#0", width=140)
        self.tree.column("status", width=60, anchor="center")
        self.tree.column("count", width=30, anchor="center")
        self.tree.pack(fill=tk.BOTH, expand=True)
        self.tree.bind("<<TreeviewSelect>>", self.on_list_select)

        # Sidebar Buttons
        btn_box = tk.Frame(sidebar, bg=self.colors['sidebar'])
        btn_box.pack(fill=tk.X, padx=10, pady=20)

        self.btn_resume = self.make_button(btn_box, "RESUME", self.resume_selected, self.colors['accent'],
                                           self.colors['accent_hover'])
        self.btn_resume.pack(fill=tk.X, pady=4)
        self.btn_resume.config(state="disabled")

        self.btn_save = self.make_button(btn_box, "SAVE & CLOSE", self.save_close_selected, self.colors['success'],
                                         self.colors['success_hover'])
        self.btn_save.pack(fill=tk.X, pady=4)
        self.btn_save.config(state="disabled")

        self.btn_discard = self.make_button(btn_box, "DISCARD", self.discard_selected, self.colors['danger'],
                                            self.colors['danger_hover'])
        self.btn_discard.pack(fill=tk.X, pady=4)
        self.btn_discard.config(state="disabled")

        # 2. Main Content Area
        main_area = tk.Frame(self.root, bg=self.colors['bg'], padx=40, pady=40)
        main_area.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        tk.Label(main_area, text="New Session", font=("Segoe UI", 24, "bold"),
                 bg=self.colors['bg'], fg=self.colors['fg']).pack(anchor="w", pady=(0, 20))

        # Form
        self.make_label(main_area, "DOCUMENT NAME")
        self.entry_name = self.make_entry(main_area)
        self.entry_name.insert(0, "screenshot")

        self.make_label(main_area, "SAVE DIRECTORY")
        dir_frame = tk.Frame(main_area, bg=self.colors['bg'])
        dir_frame.pack(fill=tk.X, pady=(0, 15))
        self.entry_dir = self.make_entry(dir_frame, side=tk.LEFT)
        self.make_small_btn(dir_frame, "...", self.browse).pack(side=tk.RIGHT, padx=(10, 0))

        self.make_label(main_area, "MODE")
        self.combo_mode = ttk.Combobox(main_area, values=["Word Document", "Folder"], state="readonly")
        self.combo_mode.current(0)
        self.combo_mode.pack(fill=tk.X, pady=(0, 20), ipady=4)

        # Options Grid
        opt_frame = tk.Frame(main_area, bg=self.colors['bg'])
        opt_frame.pack(fill=tk.X, pady=(0, 20))

        self.var_title = tk.BooleanVar()
        self.make_check(opt_frame, "Log Window Title", self.var_title).grid(row=0, column=0, sticky="w", padx=(0, 20))

        self.var_num = tk.BooleanVar(value=True)
        self.make_check(opt_frame, "Append Number", self.var_num).grid(row=0, column=1, sticky="w")

        self.var_auto = tk.BooleanVar()
        self.make_check(opt_frame, "Auto-Copy to Clipboard", self.var_auto).grid(row=1, column=0, sticky="w",
                                                                                 pady=(10, 0))

        # Copy Sub-options
        copy_sub = tk.Frame(main_area, bg=self.colors['bg'])
        copy_sub.pack(fill=tk.X, pady=(0, 20))
        self.var_copy_files = tk.BooleanVar(value=True)
        self.make_check(copy_sub, "Include Files (Explorer)", self.var_copy_files).pack(side=tk.LEFT, padx=(20, 20))
        self.var_copy_img = tk.BooleanVar(value=True)
        self.make_check(copy_sub, "Include Image (Win+V)", self.var_copy_img).pack(side=tk.LEFT)

        # Traces for mandatory selection
        self.var_auto.trace_add("write", self._validate_auto)
        self.var_copy_files.trace_add("write", self._validate_subs)
        self.var_copy_img.trace_add("write", self._validate_subs)

        # Start Button
        self.btn_start = self.make_button(main_area, "START NEW SESSION", self.start_session, self.colors['accent'],
                                          self.colors['accent_hover'])
        self.btn_start.pack(fill=tk.X, pady=(10, 0), ipady=5)
        self.btn_start.config(font=("Segoe UI", 11, "bold"))

        # Footer
        tk.Label(main_area, text="HOTKEYS:  ~ (Capture)   |   Ctrl + Alt + ~ (Undo)",
                 bg=self.colors['bg'], fg="#666666", font=("Segoe UI", 9)).pack(side=tk.BOTTOM, pady=10)

        self.lbl_status = tk.Label(main_area, text="Ready", bg=self.colors['bg'], fg=self.colors['accent'],
                                   font=("Segoe UI", 10))
        self.lbl_status.pack(side=tk.BOTTOM, pady=(0, 5))

        # --- NOTIFICATION POPUP WINDOW ---
        self.notif_window = tk.Toplevel(self.root)
        self.notif_window.withdraw()
        self.notif_window.overrideredirect(True)
        self.notif_window.attributes("-topmost", True)
        self.notif_window.configure(bg=self.colors['sidebar'])

        self.notif_frame = tk.Frame(self.notif_window, bg=self.colors['sidebar'],
                                    highlightbackground=self.colors['accent'], highlightthickness=1)
        self.notif_frame.pack(fill=tk.BOTH, expand=True)

        self.lbl_notif_title = tk.Label(self.notif_frame, text="", fg=self.colors['accent'], bg=self.colors['sidebar'],
                                        font=("Segoe UI", 10, "bold"))
        self.lbl_notif_title.pack(pady=(5, 0))

        self.lbl_notif_info = tk.Label(self.notif_frame, text="", fg=self.colors['fg'], bg=self.colors['sidebar'],
                                       font=("Segoe UI", 8))
        self.lbl_notif_info.pack(pady=(0, 5))

        # Init
        self.load_defaults()
        self.hk = HotkeyListener(self.on_hotkey_capture, self.on_hotkey_undo)
        self.hk.start()
        self.check_queue()
        self.root.protocol("WM_DELETE_WINDOW", self.on_app_close)

    # --- UI HELPERS ---
    def make_label(self, parent, text):
        tk.Label(parent, text=text, bg=self.colors['bg'], fg="#888888", font=("Segoe UI", 8, "bold")).pack(anchor="w",
                                                                                                           pady=(0, 5))

    def make_entry(self, parent, side=None):
        e = tk.Entry(parent, bg=self.colors['entry_bg'], fg=self.colors['entry_fg'],
                     insertbackground="white", relief="flat", font=("Segoe UI", 10))
        e.pack(side=side, fill=tk.X, expand=True, ipady=5) if side else e.pack(fill=tk.X, pady=(0, 15), ipady=5)
        return e

    def make_check(self, parent, text, var):
        return tk.Checkbutton(parent, text=text, variable=var, bg=self.colors['bg'], fg=self.colors['fg'],
                              selectcolor=self.colors['bg'], activebackground=self.colors['bg'],
                              activeforeground=self.colors['fg'], font=("Segoe UI", 10))

    def make_button(self, parent, text, command, bg_color, hover_color):
        btn = tk.Button(parent, text=text, command=command, bg=bg_color, fg="white",
                        disabledforeground="white", activebackground=hover_color, activeforeground="white",
                        relief="flat", borderwidth=0, font=("Segoe UI", 9, "bold"), cursor="hand2")

        def on_enter(e):
            if btn['state'] == 'normal': btn.config(bg=hover_color)

        def on_leave(e):
            if btn['state'] == 'normal': btn.config(bg=bg_color)

        btn.bind("<Enter>", on_enter)
        btn.bind("<Leave>", on_leave)
        return btn

    def make_small_btn(self, parent, text, command):
        return tk.Button(parent, text=text, command=command, bg=self.colors['entry_bg'], fg="white",
                         relief="flat", borderwidth=0, cursor="hand2", width=4)

    def show_notification(self, title, size_str):
        self.lbl_notif_title.config(text=title)
        self.lbl_notif_info.config(text=f"{size_str}")

        width, height = 180, 55
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        # Position bottom right
        x_pos = screen_width - width - 20
        y_pos = screen_height - height - 60
        self.notif_window.geometry(f"{width}x{height}+{x_pos}+{y_pos}")

        self.notif_window.deiconify()

        if self.hide_timer:
            self.root.after_cancel(self.hide_timer)
        self.hide_timer = self.root.after(1500, self.notif_window.withdraw)

    # --- FUNCTIONALITY ---
    def _validate_auto(self, *args):
        if self.var_auto.get():
            if not self.var_copy_files.get() and not self.var_copy_img.get():
                self.var_copy_img.set(True)

    def _validate_subs(self, *args):
        if self.var_auto.get():
            if not self.var_copy_files.get() and not self.var_copy_img.get():
                # Force re-check of Image if user tried to uncheck both
                self.var_copy_img.set(True)

    def browse(self):
        d = filedialog.askdirectory()
        if d:
            self.entry_dir.delete(0, tk.END)
            self.entry_dir.insert(0, d)

    def load_defaults(self):
        conf = ToonConfig.load(self.config_file)
        dd = os.path.join(os.path.expanduser("~"), "Desktop", "Evidence")
        self.entry_dir.insert(0, conf.get('save_dir', dd))
        self.entry_name.delete(0, tk.END)
        self.entry_name.insert(0, conf.get('filename', 'screenshot'))

    def save_defaults(self):
        p = self.entry_dir.get()
        if self.original_base_dir: p = self.original_base_dir
        data = {
            "filename": self.entry_name.get(),
            "save_dir": p,
            "w": self.root.winfo_width(), "h": self.root.winfo_height()
        }
        ToonConfig.save(self.config_file, data)

    def start_session(self):
        base = self.entry_dir.get().strip()

        # Determine name (Safe default)
        raw_name = self.entry_name.get().strip()
        if not raw_name: raw_name = "screenshot"

        date_str = datetime.datetime.now().strftime("%d-%m-%Y")
        if date_str not in base:
            self.original_base_dir = base
            save_dir = os.path.join(base, date_str)
            self.entry_dir.delete(0, tk.END)
            self.entry_dir.insert(0, save_dir)
        else:
            save_dir = base
            self.original_base_dir = None

        cfg = {
            "filename": raw_name,
            "save_dir": save_dir,
            "save_mode": "folder" if self.combo_mode.current() == 1 else "docx",
            "log_title": self.var_title.get(),
            "append_num": self.var_num.get(),
            "auto_copy": self.var_auto.get(),
            "copy_files": self.var_copy_files.get(),
            "copy_image": self.var_copy_img.get(),
            "max_size": "0"
        }

        sess = ScreenshotSession(cfg, self.gui_queue)

        if self.active_session_key:
            self.pause_session(self.active_session_key)

        key = sess.current_filename
        self.sessions[key] = sess
        self.active_session_key = key

        self.tree.insert("", "end", iid=key, text=os.path.basename(key), values=("Active", "0"))
        self.tree.selection_set(key)
        self.update_ui_state()

    def pause_session(self, key):
        if key in self.sessions:
            self.sessions[key].status = "Paused"
            self.tree.set(key, "status", "Paused")

    def resume_selected(self):
        sel = self.tree.selection()
        if not sel: return
        key = sel[0]

        if self.active_session_key and self.active_session_key != key:
            self.pause_session(self.active_session_key)

        self.active_session_key = key
        self.sessions[key].status = "Active"
        self.tree.set(key, "status", "Active")
        self.update_ui_state()

    def save_close_selected(self):
        sel = self.tree.selection()
        if not sel: return
        key = sel[0]
        self._close_internal(key, delete=False)

    def discard_selected(self):
        sel = self.tree.selection()
        if not sel: return
        key = sel[0]
        self._close_internal(key, delete=True)

    def _close_internal(self, key, delete):
        sess = self.sessions[key]
        sess.cleanup(delete_files=delete)

        self.tree.delete(key)
        del self.sessions[key]

        if self.active_session_key == key:
            self.active_session_key = None
            self.lbl_status.config(text="No Active Session", fg=self.colors['disabled'])

        if not self.sessions and self.original_base_dir:
            self.entry_dir.delete(0, tk.END)
            self.entry_dir.insert(0, self.original_base_dir)
            self.original_base_dir = None

    def on_list_select(self, event):
        sel = self.tree.selection()
        if not sel:
            self.btn_resume.config(state="disabled", bg=self.colors['disabled'])
            self.btn_save.config(state="disabled", bg=self.colors['disabled'])
            self.btn_discard.config(state="disabled", bg=self.colors['disabled'])
            return

        key = sel[0]
        status = self.sessions[key].status

        self.btn_save.config(state="normal", bg=self.colors['success'])
        self.btn_discard.config(state="normal", bg=self.colors['danger'])

        if status == "Paused":
            self.btn_resume.config(state="normal", bg=self.colors['accent'])
        else:
            self.btn_resume.config(state="disabled", bg=self.colors['disabled'])

    def update_ui_state(self):
        if self.active_session_key:
            name = os.path.basename(self.active_session_key)
            self.lbl_status.config(text=f"ACTIVE: {name}", fg=self.colors['accent'])

    def on_hotkey_capture(self):
        if self.active_session_key:
            self.sessions[self.active_session_key].capture()

    def on_hotkey_undo(self):
        if self.active_session_key:
            self.sessions[self.active_session_key].undo()

    def check_queue(self):
        try:
            while True:
                msg = self.gui_queue.get_nowait()
                action = msg[0]
                if action == "NOTIFY":
                    key, count, size = msg[1], msg[2], msg[3]
                    if key in self.sessions:
                        self.tree.set(key, "count", count)
                        if key == self.active_session_key:
                            self.lbl_status.config(text=f"Captured #{count} ({size})", fg=self.colors['success'])
                            # TRIGGER NOTIFICATION
                            self.show_notification(f"Saved #{count}", size)

                elif action == "UPDATE_SESSION":
                    key, count, size = msg[1], msg[2], msg[3]
                    if key in self.sessions:
                        self.tree.set(key, "count", count)
                elif action == "UNDO":
                    key, count, size = msg[1], msg[2], msg[3]
                    if key in self.sessions:
                        self.tree.set(key, "count", count)
                        if key == self.active_session_key:
                            self.lbl_status.config(text=f"Undone (#{count})", fg=self.colors['warning'])
                            # TRIGGER NOTIFICATION
                            self.show_notification(f"Undone #{count}", size)

                elif action == "WARNING":
                    messagebox.showwarning(msg[1], msg[2])
                elif action == "UPDATE_FILENAME":
                    pass
        except queue.Empty:
            pass
        self.root.after(50, self.check_queue)

    def on_app_close(self):
        if self.sessions:
            if not messagebox.askokcancel("Quit", "Open sessions will be saved. Quit?"): return
        self.save_defaults()
        self.hk.stop()
        for k, s in self.sessions.items():
            s.cleanup()
        self.root.destroy()


if __name__ == "__main__":
    ModernUI().root.mainloop()